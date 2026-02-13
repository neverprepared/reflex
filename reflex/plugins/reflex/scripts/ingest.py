#!/usr/bin/env python3
"""
Ingest local files into Qdrant vector database.

Supported formats:
- PDF (.pdf)
- Markdown (.md)
- Text (.txt, .rst)
- HTML (.html, .htm)
- EPUB (.epub)
- Word (.docx)
- Jupyter notebooks (.ipynb)
- Mermaid diagrams (.mmd, .mermaid)
- Code files (.py, .js, .ts, .go, .rs, .java, .c, .cpp, .rb, .sh)

Usage:
    uvx --with pymupdf,fastembed,qdrant-client,python-docx,ebooklib,beautifulsoup4 \
        python ingest.py <path> [--collection NAME] [--chunk-size WORDS]

Examples:
    python ingest.py ~/Documents/manual.pdf
    python ingest.py ~/notes/ --collection research
    python ingest.py ~/code/project --chunk-size 300
"""

import argparse
import hashlib
import json
import os
import re
import sys
import time
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple


# =============================================================================
# Custom Exceptions
# =============================================================================

class IngestError(Exception):
    """Base exception for ingestion errors."""
    pass


class DependencyError(IngestError):
    """Raised when a required dependency is missing."""
    pass


class ExtractorError(IngestError):
    """Raised when text extraction fails."""
    pass


class QdrantConnectionError(IngestError):
    """Raised when Qdrant connection fails."""
    pass


class FileSizeError(IngestError):
    """Raised when file exceeds size limit."""
    pass


# =============================================================================
# Constants
# =============================================================================

MAX_FILE_SIZE = 100 * 1024 * 1024  # 100MB


# =============================================================================
# Dependency imports with helpful error messages
# =============================================================================

try:
    from qdrant_client import QdrantClient
    from qdrant_client.models import PointStruct, VectorParams, Distance
except ImportError:
    raise DependencyError(
        "qdrant-client not installed. Run: pip install qdrant-client"
    )

try:
    from fastembed import TextEmbedding
except ImportError:
    raise DependencyError(
        "fastembed not installed. Run: pip install fastembed"
    )


# =============================================================================
# Extractors for different file formats
# =============================================================================

def extract_pdf(path: Path) -> Tuple[str, Dict]:
    """Extract text from PDF using PyMuPDF."""
    try:
        import fitz
    except ImportError:
        raise ExtractorError("pymupdf not installed. Run: pip install pymupdf")

    doc = fitz.open(str(path))
    pages = []
    for page in doc:
        text = page.get_text().strip()
        if text:
            pages.append(text)
    doc.close()

    return "\n\n".join(pages), {"total_pages": len(pages), "format": "pdf"}


def extract_markdown(path: Path) -> Tuple[str, Dict]:
    """Extract text from Markdown file."""
    text = path.read_text(encoding="utf-8", errors="ignore")
    return text, {"format": "markdown"}


def extract_text(path: Path) -> Tuple[str, Dict]:
    """Extract text from plain text file."""
    text = path.read_text(encoding="utf-8", errors="ignore")
    return text, {"format": "text"}


def extract_html(path: Path) -> Tuple[str, Dict]:
    """Extract text from HTML file."""
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        raise ExtractorError("beautifulsoup4 not installed. Run: pip install beautifulsoup4")

    html = path.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(html, "html.parser")

    # Remove scripts, styles, nav, footer
    for tag in soup.find_all(["script", "style", "nav", "footer", "aside"]):
        tag.decompose()

    # Get title
    title = soup.title.string if soup.title else path.stem

    # Get main content
    main = soup.find("main") or soup.find("article") or soup.find("body")
    text = main.get_text(separator="\n", strip=True) if main else ""

    return text, {"format": "html", "title": title}


def extract_epub(path: Path) -> Tuple[str, Dict]:
    """Extract text from EPUB ebook."""
    try:
        import ebooklib
        from ebooklib import epub
        from bs4 import BeautifulSoup
    except ImportError:
        raise ExtractorError("ebooklib not installed. Run: pip install ebooklib beautifulsoup4")

    book = epub.read_epub(str(path))

    # Get metadata
    title = book.get_metadata("DC", "title")
    title = title[0][0] if title else path.stem

    author = book.get_metadata("DC", "creator")
    author = author[0][0] if author else None

    # Extract text from all documents
    chapters = []
    for item in book.get_items():
        if item.get_type() == ebooklib.ITEM_DOCUMENT:
            soup = BeautifulSoup(item.get_content(), "html.parser")
            text = soup.get_text(separator="\n", strip=True)
            if text:
                chapters.append(text)

    return "\n\n".join(chapters), {
        "format": "epub",
        "title": title,
        "author": author,
        "chapters": len(chapters)
    }


def extract_docx(path: Path) -> Tuple[str, Dict]:
    """Extract text from Word document."""
    try:
        from docx import Document
    except ImportError:
        raise ExtractorError("python-docx not installed. Run: pip install python-docx")

    doc = Document(str(path))

    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.replace("|", "").strip():
                paragraphs.append(row_text)

    return "\n\n".join(paragraphs), {"format": "docx"}


def extract_notebook(path: Path) -> Tuple[str, Dict]:
    """Extract text from Jupyter notebook."""
    content = json.loads(path.read_text(encoding="utf-8"))

    cells = []
    code_cells = 0
    markdown_cells = 0

    for cell in content.get("cells", []):
        cell_type = cell.get("cell_type", "")
        source = "".join(cell.get("source", []))

        if cell_type == "markdown":
            cells.append(source)
            markdown_cells += 1
        elif cell_type == "code":
            cells.append(f"```python\n{source}\n```")
            code_cells += 1

    return "\n\n".join(cells), {
        "format": "jupyter",
        "code_cells": code_cells,
        "markdown_cells": markdown_cells
    }


def extract_mermaid(path: Path) -> Tuple[str, Dict]:
    """Extract Mermaid diagram with metadata."""
    text = path.read_text(encoding="utf-8", errors="ignore")

    # Detect diagram type from content
    diagram_type = detect_mermaid_type(text)

    # Detect architectural patterns
    patterns = detect_diagram_patterns(text)

    # Extract components (node names)
    components = extract_mermaid_components(text)

    # Look for description in comments or preceding text
    description = extract_mermaid_description(text)

    # Build rich document for embedding
    if description:
        document = f"{description}\n\n```mermaid\n{text}\n```"
    else:
        document = f"Mermaid {diagram_type} diagram\n\nComponents: {', '.join(components[:10])}\n\n```mermaid\n{text}\n```"

    return document, {
        "format": "mermaid",
        "diagram_type": diagram_type,
        "patterns": patterns,
        "components": components[:20],  # Limit stored components
    }


def detect_mermaid_type(text: str) -> str:
    """Detect the type of Mermaid diagram."""
    # Skip comment lines to find the diagram declaration
    lines = text.strip().split('\n')
    for line in lines:
        line = line.strip().lower()
        if line.startswith('%%') or not line:
            continue
        text_lower = line
        break
    else:
        text_lower = text.strip().lower()

    if text_lower.startswith("graph ") or text_lower.startswith("flowchart "):
        return "flowchart"
    elif text_lower.startswith("sequencediagram"):
        return "sequence"
    elif text_lower.startswith("classdiagram"):
        return "class"
    elif text_lower.startswith("statediagram"):
        return "state"
    elif text_lower.startswith("erdiagram"):
        return "er"
    elif text_lower.startswith("gantt"):
        return "gantt"
    elif text_lower.startswith("pie"):
        return "pie"
    elif text_lower.startswith("journey"):
        return "journey"
    elif text_lower.startswith("gitgraph"):
        return "git"
    elif text_lower.startswith("c4context") or text_lower.startswith("c4container"):
        return "c4"
    elif text_lower.startswith("mindmap"):
        return "mindmap"
    elif text_lower.startswith("timeline"):
        return "timeline"
    elif text_lower.startswith("architecture"):
        return "architecture"
    else:
        return "unknown"


def detect_diagram_patterns(text: str) -> List[str]:
    """Detect architectural patterns in diagram."""
    patterns = []
    text_lower = text.lower()

    # Common architectural patterns
    pattern_keywords = {
        "microservices": ["service", "api", "gateway", "mesh"],
        "event-driven": ["event", "queue", "publish", "subscribe", "kafka", "rabbitmq"],
        "layered": ["presentation", "business", "data", "layer", "controller", "service", "repository"],
        "client-server": ["client", "server", "request", "response"],
        "mvc": ["model", "view", "controller"],
        "cqrs": ["command", "query", "read", "write"],
        "saga": ["saga", "orchestrat", "compensat"],
        "circuit-breaker": ["circuit", "breaker", "fallback", "retry"],
        "api-gateway": ["gateway", "api", "route", "proxy"],
        "database": ["database", "db", "postgres", "mysql", "mongo", "redis"],
        "authentication": ["auth", "login", "token", "jwt", "oauth", "sso"],
        "caching": ["cache", "redis", "memcache"],
        "load-balancing": ["load", "balancer", "nginx", "haproxy"],
        "pub-sub": ["pub", "sub", "topic", "subscriber", "publisher"],
    }

    for pattern, keywords in pattern_keywords.items():
        if any(kw in text_lower for kw in keywords):
            patterns.append(pattern)

    return patterns


def extract_mermaid_components(text: str) -> List[str]:
    """Extract component/node names from Mermaid diagram."""
    components = set()

    # Match node definitions: A[Label], B(Label), C{Label}, D((Label))
    node_patterns = [
        r'(\w+)\[([^\]]+)\]',      # A[Label]
        r'(\w+)\(([^)]+)\)',        # B(Label)
        r'(\w+)\{([^}]+)\}',        # C{Label}
        r'(\w+)\[\[([^\]]+)\]\]',   # D[[Label]]
        r'(\w+)\(\(([^)]+)\)\)',    # E((Label))
        r'(\w+)>([^]]+)\]',         # F>Label]
    ]

    for pattern in node_patterns:
        for match in re.finditer(pattern, text):
            # Add both the ID and the label
            components.add(match.group(1))
            label = match.group(2).strip()
            if label and len(label) < 50:  # Reasonable label length
                components.add(label)

    # Match participant declarations in sequence diagrams
    participant_match = re.findall(r'participant\s+(\w+)', text, re.IGNORECASE)
    components.update(participant_match)

    # Match actor declarations
    actor_match = re.findall(r'actor\s+(\w+)', text, re.IGNORECASE)
    components.update(actor_match)

    return list(components)


def extract_mermaid_description(text: str) -> Optional[str]:
    """Extract description from Mermaid file comments."""
    lines = text.split('\n')
    description_lines = []

    for line in lines:
        # Mermaid comments start with %%
        if line.strip().startswith('%%'):
            comment = line.strip()[2:].strip()
            if comment and not comment.startswith('{'):  # Skip directives
                description_lines.append(comment)
        elif line.strip() and not line.strip().startswith('%%'):
            # Stop at first non-comment line
            break

    return ' '.join(description_lines) if description_lines else None


def extract_code(path: Path) -> Tuple[str, Dict]:
    """Extract text from code file."""
    text = path.read_text(encoding="utf-8", errors="ignore")

    # Detect language from extension
    ext_to_lang = {
        ".py": "python",
        ".js": "javascript",
        ".ts": "typescript",
        ".go": "go",
        ".rs": "rust",
        ".java": "java",
        ".c": "c",
        ".cpp": "cpp",
        ".h": "c",
        ".hpp": "cpp",
        ".rb": "ruby",
        ".sh": "bash",
        ".bash": "bash",
        ".zsh": "zsh",
        ".sql": "sql",
        ".yaml": "yaml",
        ".yml": "yaml",
        ".json": "json",
        ".toml": "toml",
    }

    lang = ext_to_lang.get(path.suffix.lower(), "text")

    return text, {"format": "code", "language": lang}


# =============================================================================
# Format detection and routing
# =============================================================================

EXTRACTORS = {
    ".pdf": extract_pdf,
    ".md": extract_markdown,
    ".markdown": extract_markdown,
    ".txt": extract_text,
    ".rst": extract_text,
    ".html": extract_html,
    ".htm": extract_html,
    ".epub": extract_epub,
    ".docx": extract_docx,
    ".ipynb": extract_notebook,
    # Diagrams
    ".mmd": extract_mermaid,
    ".mermaid": extract_mermaid,
    # Code files
    ".py": extract_code,
    ".js": extract_code,
    ".ts": extract_code,
    ".go": extract_code,
    ".rs": extract_code,
    ".java": extract_code,
    ".c": extract_code,
    ".cpp": extract_code,
    ".h": extract_code,
    ".hpp": extract_code,
    ".rb": extract_code,
    ".sh": extract_code,
    ".bash": extract_code,
    ".sql": extract_code,
    ".yaml": extract_code,
    ".yml": extract_code,
}


def is_supported(path: Path) -> bool:
    """Check if file format is supported."""
    return path.suffix.lower() in EXTRACTORS


def extract(path: Path) -> Tuple[str, Dict]:
    """Extract text from file based on extension."""
    ext = path.suffix.lower()
    extractor = EXTRACTORS.get(ext)

    if not extractor:
        raise ValueError(f"Unsupported format: {ext}")

    return extractor(path)


# =============================================================================
# Chunking
# =============================================================================

def chunk_text(
    text: str,
    chunk_size: int = 400,
    overlap: int = 50
) -> List[Dict]:
    """
    Chunk text into overlapping segments.

    Args:
        text: Full text to chunk
        chunk_size: Target words per chunk
        overlap: Words to overlap between chunks

    Returns:
        List of chunk dicts with content and word_count
    """
    # Split into paragraphs
    paragraphs = [p.strip() for p in re.split(r'\n\s*\n', text) if p.strip()]

    chunks = []
    current = []
    current_words = 0

    for para in paragraphs:
        para_words = len(para.split())

        # If adding this paragraph exceeds limit, save current chunk
        if current_words + para_words > chunk_size and current:
            chunks.append({
                "content": "\n\n".join(current),
                "word_count": current_words
            })

            # Keep last paragraph for overlap
            if current and len(current[-1].split()) <= overlap:
                current = [current[-1]]
                current_words = len(current[-1].split())
            else:
                current = []
                current_words = 0

        current.append(para)
        current_words += para_words

    # Don't forget last chunk
    if current:
        chunks.append({
            "content": "\n\n".join(current),
            "word_count": current_words
        })

    return chunks


# =============================================================================
# Qdrant ingestion
# =============================================================================

def connect_to_qdrant(
    qdrant_url: str,
    max_retries: int = 3,
    retry_delay: float = 2.0
) -> QdrantClient:
    """
    Connect to Qdrant with retry logic.

    Args:
        qdrant_url: Qdrant server URL
        max_retries: Maximum number of connection attempts
        retry_delay: Seconds to wait between retries

    Returns:
        Connected QdrantClient

    Raises:
        QdrantConnectionError: If all connection attempts fail
    """
    last_error = None

    for attempt in range(1, max_retries + 1):
        try:
            client = QdrantClient(url=qdrant_url, timeout=10)
            # Test connection by listing collections
            client.get_collections()
            return client
        except Exception as e:
            last_error = e
            if attempt < max_retries:
                print(f"  Connection attempt {attempt}/{max_retries} failed: {e}")
                print(f"  Retrying in {retry_delay}s...")
                time.sleep(retry_delay)
            else:
                raise QdrantConnectionError(
                    f"Failed to connect to Qdrant at {qdrant_url} after {max_retries} attempts: {last_error}"
                )

    # Should not reach here, but just in case
    raise QdrantConnectionError(f"Failed to connect to Qdrant: {last_error}")


def ingest_to_qdrant(
    chunks: List[Dict],
    file_path: Path,
    file_metadata: Dict,
    collection: str,
    qdrant_url: str = "http://localhost:6333"
) -> int:
    """
    Ingest chunks into Qdrant.

    Args:
        chunks: List of chunk dicts
        file_path: Original file path
        file_metadata: Metadata from extraction
        collection: Qdrant collection name
        qdrant_url: Qdrant server URL

    Returns:
        Number of chunks ingested

    Raises:
        QdrantConnectionError: If connection to Qdrant fails
    """
    # Initialize clients with retry
    client = connect_to_qdrant(qdrant_url)
    embedder = TextEmbedding("sentence-transformers/all-MiniLM-L6-v2")

    # Vector name used by mcp-server-qdrant
    vector_name = "fast-all-minilm-l6-v2"

    # Ensure collection exists
    collections = [c.name for c in client.get_collections().collections]
    if collection not in collections:
        client.create_collection(
            collection_name=collection,
            vectors_config={
                vector_name: VectorParams(size=384, distance=Distance.COSINE)
            }
        )
        print(f"Created collection: {collection}")

    # Prepare points
    points = []
    contents = [c["content"] for c in chunks]

    print(f"Generating embeddings for {len(chunks)} chunks...")
    embeddings = list(embedder.embed(contents))

    # Generate file hash for deduplication
    file_hash = hashlib.md5(file_path.read_bytes()).hexdigest()[:12]

    for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
        # Create unique ID from file hash + chunk index
        point_id = hashlib.md5(f"{file_hash}_{i}".encode()).hexdigest()

        # Build metadata
        metadata = {
            "source": "local_file",
            "content_type": file_metadata.get("format", "text"),
            "original_path": str(file_path.absolute()),
            "filename": file_path.name,
            "harvested_at": datetime.now().isoformat(),
            "chunk_index": i,
            "total_chunks": len(chunks),
            "word_count": chunk["word_count"],
            **{k: v for k, v in file_metadata.items() if v is not None}
        }

        # Use named vector to match mcp-server-qdrant format
        points.append(PointStruct(
            id=point_id,
            vector={vector_name: embedding.tolist()},
            payload={
                "document": chunk["content"],
                "metadata": metadata
            }
        ))

    # Upsert in batches
    batch_size = 100
    for i in range(0, len(points), batch_size):
        batch = points[i:i + batch_size]
        client.upsert(collection_name=collection, points=batch)
        print(f"  Ingested {min(i + batch_size, len(points))}/{len(points)} chunks")

    return len(points)


# =============================================================================
# Main
# =============================================================================

def process_file(
    path: Path,
    collection: str,
    chunk_size: int,
    qdrant_url: str,
    max_file_size: int = MAX_FILE_SIZE
) -> Dict:
    """Process a single file."""
    print(f"\nProcessing: {path.name}")

    # Check file size
    file_size = path.stat().st_size
    if file_size > max_file_size:
        size_mb = file_size / (1024 * 1024)
        limit_mb = max_file_size / (1024 * 1024)
        raise FileSizeError(
            f"File size ({size_mb:.1f}MB) exceeds limit ({limit_mb:.0f}MB)"
        )

    # Extract text
    text, metadata = extract(path)

    if not text.strip():
        print(f"  Warning: No text extracted from {path.name}")
        return {"file": str(path), "status": "empty", "chunks": 0}

    # Chunk
    chunks = chunk_text(text, chunk_size=chunk_size)
    print(f"  Extracted {len(text.split())} words -> {len(chunks)} chunks")

    # Ingest
    ingested = ingest_to_qdrant(
        chunks=chunks,
        file_path=path,
        file_metadata=metadata,
        collection=collection,
        qdrant_url=qdrant_url
    )

    return {
        "file": str(path),
        "status": "success",
        "chunks": ingested,
        "format": metadata.get("format", "unknown")
    }


def main():
    parser = argparse.ArgumentParser(
        description="Ingest local files into Qdrant vector database"
    )
    parser.add_argument(
        "path",
        type=Path,
        help="File or directory to ingest"
    )
    parser.add_argument(
        "--collection",
        default=os.environ.get("COLLECTION_NAME", "personal_memories"),
        help="Qdrant collection name (default: personal_memories)"
    )
    parser.add_argument(
        "--chunk-size",
        type=int,
        default=400,
        help="Target words per chunk (default: 400)"
    )
    parser.add_argument(
        "--qdrant-url",
        default=os.environ.get("QDRANT_URL", "http://localhost:6333"),
        help="Qdrant server URL (default: http://localhost:6333)"
    )
    parser.add_argument(
        "--recursive",
        action="store_true",
        help="Recursively process directories"
    )
    parser.add_argument(
        "--max-file-size",
        type=int,
        default=int(os.environ.get("MAX_FILE_SIZE_MB", 100)),
        help="Maximum file size in MB (default: 100, or MAX_FILE_SIZE_MB env var)"
    )

    args = parser.parse_args()

    # Convert MB to bytes for file size limit
    max_file_size = args.max_file_size * 1024 * 1024

    # Collect files to process
    files = []
    if args.path.is_file():
        if is_supported(args.path):
            files.append(args.path)
        else:
            print(f"Unsupported format: {args.path.suffix}")
            sys.exit(1)
    elif args.path.is_dir():
        pattern = "**/*" if args.recursive else "*"
        for p in args.path.glob(pattern):
            if p.is_file() and is_supported(p):
                files.append(p)

        if not files:
            print(f"No supported files found in {args.path}")
            sys.exit(1)
    else:
        print(f"Path not found: {args.path}")
        sys.exit(1)

    print(f"Found {len(files)} file(s) to ingest")
    print(f"Collection: {args.collection}")
    print(f"Chunk size: {args.chunk_size} words")

    # Process files
    results = []
    for path in files:
        try:
            result = process_file(
                path=path,
                collection=args.collection,
                chunk_size=args.chunk_size,
                qdrant_url=args.qdrant_url,
                max_file_size=max_file_size
            )
            results.append(result)
        except QdrantConnectionError as e:
            # Connection errors are fatal - stop processing
            print(f"\nFatal: {e}")
            sys.exit(1)
        except FileSizeError as e:
            # File too large - skip with warning
            print(f"  Skipped: {e}")
            results.append({
                "file": str(path),
                "status": "skipped",
                "error": str(e)
            })
        except (ExtractorError, IngestError) as e:
            # Extraction/ingestion errors - log and continue
            print(f"  Error: {e}")
            results.append({
                "file": str(path),
                "status": "error",
                "error": str(e)
            })
        except Exception as e:
            # Unexpected errors - log with details and continue
            print(f"  Unexpected error: {type(e).__name__}: {e}")
            results.append({
                "file": str(path),
                "status": "error",
                "error": f"{type(e).__name__}: {e}"
            })

    # Summary
    print("\n" + "=" * 50)
    print("Summary:")
    total_chunks = sum(r.get("chunks", 0) for r in results)
    success = sum(1 for r in results if r["status"] == "success")
    skipped = sum(1 for r in results if r["status"] == "skipped")
    errors = sum(1 for r in results if r["status"] == "error")

    print(f"  Files processed: {len(results)}")
    print(f"  Successful: {success}")
    if skipped:
        print(f"  Skipped (too large): {skipped}")
    print(f"  Errors: {errors}")
    print(f"  Total chunks ingested: {total_chunks}")


if __name__ == "__main__":
    main()
